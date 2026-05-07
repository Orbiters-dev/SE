"""
read_docx.py
DOCX 파일 내용을 텍스트로 출력
사용법: python tools/read_docx.py path/to/file.docx
"""
import sys
from pathlib import Path

def read_docx(path: str):
    try:
        from docx import Document
    except ImportError:
        print("[ERROR] pip install python-docx")
        sys.exit(1)

    doc = Document(path)
    for para in doc.paragraphs:
        print(para.text)

    # 표 내용도 출력
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            print(" | ".join(cells))

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("사용법: python tools/read_docx.py 파일경로.docx")
        sys.exit(1)
    read_docx(sys.argv[1])

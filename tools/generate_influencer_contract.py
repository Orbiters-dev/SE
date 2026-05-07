"""
generate_influencer_contract.py
인플루언서 협업 계약서 생성기

원본 DOCX 템플릿 기반 (서식 그대로 유지):
  인플루언서 컨택/형식/インフルエンサーコンテンツ契約書（有償） 5.docx
  인플루언서 컨택/형식/インフルエンサーコンテンツ契約書（無償） 4.docx

사용법:
  python tools/generate_influencer_contract.py --manual '{"collab_type":"gifting","influencer_name":"田中さくら","influencer_handle":"@sakura_life","influencer_email":"sakura@email.com","products_description":"PPSU 240ml ストローマグ 2個"}'
"""

import argparse
import json
import os
import sys
import uuid
from datetime import date, datetime
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")
sys.stderr.reconfigure(encoding="utf-8")


# 원본 DOCX 템플릿 경로 (세은 관리 원본 — 참고자료)
# 2026-04-22 스위치: 숫자 없는 깨끗한 버전 ({{VAR}} 플레이스홀더 + _____ AUTODETECT용)
TEMPLATE_DIR = Path(r"C:\Users\orbit\Desktop\s\참고자료\인플루언서 컨택\형식")
TEMPLATE_PAID    = TEMPLATE_DIR / "インフルエンサーコンテンツ契約書（有償）.docx"
TEMPLATE_GIFTING = TEMPLATE_DIR / "インフルエンサーコンテンツ契約書（無償）.docx"


def check_deps():
    try:
        import docx  # noqa
    except ImportError:
        print("[ERROR] python-docx 없음. 설치: pip install python-docx")
        sys.exit(1)


def load_negotiation(handle: str) -> dict:
    handle_clean = handle.lstrip("@")
    path = Path(".tmp/influencer_negotiations") / f"{handle_clean}.json"
    if not path.exists():
        print(f"[ERROR] 협상 데이터 없음: {path}")
        print("  --manual 옵션으로 직접 입력하거나 협상 기록을 먼저 저장하세요.")
        sys.exit(1)
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def load_manual(json_str: str) -> dict:
    try:
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        print(f"[ERROR] JSON 파싱 실패: {e}")
        sys.exit(1)


def validate_fields(data: dict) -> dict:
    """필수 필드 검증 + 기본값"""
    collab_type = data.get("collab_type", "").lower()
    if collab_type not in ("paid", "gifting"):
        print(f"[ERROR] collab_type은 'paid' 또는 'gifting'이어야 합니다. 현재: '{collab_type}'")
        sys.exit(1)

    required_common = ["influencer_name", "influencer_handle", "deliverables_detail", "product_type"]
    required_paid = ["compensation_amount"]

    missing = [f for f in required_common if not data.get(f)]
    if collab_type == "paid":
        missing += [f for f in required_paid if not data.get(f)]

    if missing:
        print(f"[ERROR] 필수 필드 누락: {', '.join(missing)}")
        print(f"  - deliverables_detail 예: 'Instagram Reels 1本'")
        print(f"  - product_type: 'ppsu_straw' / 'ppsu_onetouch' / 'stainless' 중 하나")
        sys.exit(1)

    # product_type 검증
    valid_product_types = ("ppsu_straw", "ppsu_onetouch", "stainless")
    if data.get("product_type") not in valid_product_types:
        print(f"[ERROR] product_type 값 오류: '{data.get('product_type')}'")
        print(f"  → 허용값: {', '.join(valid_product_types)}")
        sys.exit(1)

    today = date.today()
    data.setdefault("contract_date", today.strftime("%Y年%m月%d日"))
    data.setdefault("products_description", "")
    data.setdefault("compensation_currency", "JPY")
    data.setdefault("whitelisting_days", 90)
    data.setdefault("contract_id", f"IC-{today.strftime('%Y%m')}-{str(uuid.uuid4())[:6].upper()}")

    # handle에서 @ 제거
    data["influencer_handle"] = data["influencer_handle"].lstrip("@")

    return data


def replace_text_in_paragraph(para, old: str, new: str, bold_segment: str = None):
    """
    단락 내 텍스트 치환. runs에 걸쳐 분리된 경우도 처리.
    서식(볼드, 폰트 등)은 첫 번째 run 기준 유지.
    치환 후 빨간색 폰트는 검정으로 리셋.
    bold_segment가 주어지면 해당 부분만 볼드 run으로 분리.
    """
    from docx.shared import RGBColor
    from copy import deepcopy
    full_text = para.text
    if old not in full_text:
        return

    new_full = full_text.replace(old, new)

    if para.runs:
        # 첫 run의 서식 정보 보존
        base_run = para.runs[0]
        base_font = base_run.font

        # 모든 run 텍스트 비움 (첫 run 포함)
        for run in para.runs:
            run.text = ""

        if bold_segment and bold_segment in new_full:
            # bold_segment 기준으로 분리: before + bold + after
            idx = new_full.index(bold_segment)
            before = new_full[:idx]
            after = new_full[idx + len(bold_segment):]

            base_run.text = before
            base_run.font.color.rgb = RGBColor(0, 0, 0)

            # 볼드 run 추가
            bold_run = para.add_run(bold_segment)
            bold_run.bold = True
            bold_run.font.color.rgb = RGBColor(0, 0, 0)
            if base_font.name:
                bold_run.font.name = base_font.name
            if base_font.size:
                bold_run.font.size = base_font.size

            # after run 추가
            if after:
                after_run = para.add_run(after)
                after_run.font.color.rgb = RGBColor(0, 0, 0)
                if base_font.name:
                    after_run.font.name = base_font.name
                if base_font.size:
                    after_run.font.size = base_font.size
        else:
            base_run.text = new_full
            base_run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        if bold_segment and bold_segment in new_full:
            idx = new_full.index(bold_segment)
            before = new_full[:idx]
            after = new_full[idx + len(bold_segment):]
            if before:
                r = para.add_run(before)
                r.font.color.rgb = RGBColor(0, 0, 0)
            bold_run = para.add_run(bold_segment)
            bold_run.bold = True
            bold_run.font.color.rgb = RGBColor(0, 0, 0)
            if after:
                r = para.add_run(after)
                r.font.color.rgb = RGBColor(0, 0, 0)
        else:
            r = para.add_run(new_full)
            r.font.color.rgb = RGBColor(0, 0, 0)


def build_docx_from_template(template_path: Path, output_path: Path, data: dict):
    """원본 DOCX 복사 후 플레이스홀더 치환"""
    from docx import Document

    if not template_path.exists():
        print(f"[ERROR] 템플릿 없음: {template_path}")
        sys.exit(1)

    doc = Document(str(template_path))

    collab_type = data["collab_type"]
    handle = data["influencer_handle"]
    name = data["influencer_name"]
    contract_date = data["contract_date"]

    # IG 핸들 검증: 영문소문자/숫자/언더스코어/피리어드만 허용
    import re as _re_handle
    clean_handle = handle.lstrip("@")
    if not _re_handle.match(r'^[a-z0-9_.]+$', clean_handle):
        print(f"[ERROR] IG 핸들이 올바르지 않습니다: {handle}")
        print(f"  → 영문소문자/숫자/언더스코어/피리어드만 가능합니다.")
        print(f"  → 일본어 닉네임이나 호칭을 핸들로 넣지 마세요.")
        print(f"  → 세은에게 정확한 핸들을 확인하세요.")
        sys.exit(1)

    # 치환 목록
    # \u2028 = LINE SEPARATOR — 甲 서명란 단락에만 붙어 있어서 乙과 구분 가능
    COMPANY_SIGNER_NAME = "許世恩"

    # deliverables_detail: 협의된 납품물 명세 (예: "Instagram Reels 1本・YouTube Shorts 1本")
    deliverables_detail = data.get(
        "deliverables_detail",
        "TikTok／Instagram Reels／YouTube Shorts いずれか"
    )

    # 제품별 해시태그 맵 (2026-04-21 세은 통일 룰 적용)
    # - 공통 베이스: #グロミミ #grosmimi #ストローマグ #スマートマグ
    # - PPSU / ワンタッチ는 #PPSU 추가
    # - 스텐레스는 베이스 그대로
    # - 유상/무상 구분 없음 (#PR 제거)
    HASHTAG_MAP_GIFTING = {
        "ppsu_straw":    "#グロミミ #grosmimi #ストローマグ #スマートマグ #PPSU",
        "ppsu_onetouch": "#グロミミ #grosmimi #ストローマグ #スマートマグ #PPSU",
        "stainless":     "#グロミミ #grosmimi #ストローマグ #スマートマグ",
    }
    HASHTAG_MAP_PAID = {
        "ppsu_straw":    "#グロミミ #grosmimi #ストローマグ #スマートマグ #PPSU",
        "ppsu_onetouch": "#グロミミ #grosmimi #ストローマグ #スマートマグ #PPSU",
        "stainless":     "#グロミミ #grosmimi #ストローマグ #スマートマグ",
    }
    product_type = data.get("product_type", "")
    hashtag_map = HASHTAG_MAP_PAID if collab_type == "paid" else HASHTAG_MAP_GIFTING
    hashtags = hashtag_map.get(product_type, "")

    # 2026-04-22 스위치: 참고자료 템플릿 {{VAR}} 기반 치환
    # 납품물 포맷: "(Instagram Reels)" 같은 괄호 포함 플랫폼 문자열
    platforms_fmt = f"（{deliverables_detail}）"

    replacements = {
        "{{AGREEMENT_DATE}}": contract_date,
        "{{ACCOUNT_HANDLE}}": handle,
        "{{INFLUENCER_NAME}}": name,
        "{{PLATFORMS}}": platforms_fmt,
        "{{COMPANY_NAME}}": COMPANY_SIGNER_NAME,
    }

    # 해시태그 치환 ({{HASHTAGS}} 플레이스홀더로 통일)
    if hashtags:
        replacements["{{HASHTAGS}}"] = hashtags

    # bold_map: 특정 치환에서 볼드 처리할 부분 지정
    bold_map = {}
    bold_map["{{PLATFORMS}}"] = platforms_fmt
    if hashtags:
        bold_map["{{HASHTAGS}}"] = hashtags

    if collab_type == "paid":
        amount_raw = data.get("compensation_amount", "")
        # 숫자면 쉼표 포맷, 아니면 그대로
        try:
            amount_fmt = f"{int(str(amount_raw).replace(',', '')):,}"
        except (ValueError, TypeError):
            amount_fmt = str(amount_raw)
        replacements["{{PAYMENT_AMOUNT}}"] = amount_fmt
        bold_map["{{PAYMENT_AMOUNT}}"] = amount_fmt

    # 본문 단락 치환
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                replace_text_in_paragraph(para, old, new, bold_segment=bold_map.get(old))

    # 테이블 내 치환
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in para.text:
                            replace_text_in_paragraph(para, old, new, bold_segment=bold_map.get(old))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"  저장: {output_path}")


def save_record(data: dict, files: list, output_dir: Path) -> Path:
    collab_type = data["collab_type"]
    record = {
        "contract_id": data["contract_id"],
        "generated_at": datetime.now().isoformat(),
        "collab_type": collab_type,
        "influencer": {
            "name": data["influencer_name"],
            "handle": "@" + data["influencer_handle"],
            "email": data.get("influencer_email", ""),
        },
        "terms": {
            "products_description": data.get("products_description", ""),
            "whitelisting_days": data.get("whitelisting_days", 90),
            **({"compensation": f"{data['compensation_amount']} {data.get('compensation_currency','JPY')}"} if collab_type == "paid" else {}),
        },
        "files": [str(f) for f in files],
        "status": "generated",
        "docusign_sent": False,
        "docusign_envelope_id": None,
        "signed_at": None,
    }
    record_path = output_dir / f"contract_{data['influencer_handle']}_record.json"
    with open(record_path, "w", encoding="utf-8") as f:
        json.dump(record, f, ensure_ascii=False, indent=2)
    print(f"  기록: {record_path}")
    return record_path


def print_summary(data: dict, files: list):
    collab_type = data["collab_type"]
    print("\n" + "=" * 60)
    print(f"계약서 생성 완료  [{collab_type.upper()}]")
    print("=" * 60)
    print(f"  계약 ID    : {data['contract_id']}")
    print(f"  인플루언서 : {data['influencer_name']} (@{data['influencer_handle']})")
    if collab_type == "paid":
        print(f"  보상       : {data['compensation_amount']} {data.get('compensation_currency','JPY')}")
    else:
        print(f"  보상       : 기프팅 제품 (무상)")
    print("\n  생성 파일:")
    for f in files:
        print(f"    {f}")
    print("\n  다음 단계:")
    print("  1. DOCX 열어서 내용 최종 확인")
    print("  2. DocuSign 업로드: python tools/send_docusign_contract.py --handle @핸들")
    print("=" * 60)


def main():
    parser = argparse.ArgumentParser(description="인플루언서 협업 계약서 생성 (원본 DOCX 템플릿 기반)")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--handle", help="인스타 핸들 (협상 JSON 자동 로드)")
    group.add_argument("--manual", help="직접 JSON 입력")
    parser.add_argument("--output", default=r"C:\Users\orbit\Desktop\s\인플루언서 계약서(서명 없는 ver)",
                        help="출력 디렉토리")
    args = parser.parse_args()

    check_deps()

    data = load_negotiation(args.handle) if args.handle else load_manual(args.manual)
    data = validate_fields(data)

    collab_type = data["collab_type"]
    handle_clean = data["influencer_handle"]
    influencer_name = data["influencer_name"]
    date_str = date.today().strftime("%Y%m%d")
    output_dir = Path(args.output)

    template_path = TEMPLATE_PAID if collab_type == "paid" else TEMPLATE_GIFTING

    print(f"\n[계약서 생성]")
    print(f"  유형     : {collab_type.upper()}")
    print(f"  대상     : {influencer_name} (@{handle_clean})")
    print(f"  템플릿   : {template_path.name}")

    filename_base = f"{influencer_name}_{date_str}"
    docx_path = output_dir / f"{filename_base}.docx"

    build_docx_from_template(template_path, docx_path, data)

    # PDF 자동 생성
    pdf_path = output_dir / f"{filename_base}.pdf"
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        print(f"  저장: {pdf_path}")
    except ImportError:
        print("[ERROR] docx2pdf 없음. 설치: pip install docx2pdf")
        pdf_path = None
    except Exception as e:
        print(f"[ERROR] PDF 변환 실패: {e}")
        pdf_path = None

    files = [docx_path] + ([pdf_path] if pdf_path else [])

    # record는 기존 경로에 저장 (메타데이터)
    record_dir = Path("Data Storage/contracts/influencer")
    save_record(data, files, record_dir)
    print_summary(data, files)


if __name__ == "__main__":
    main()
